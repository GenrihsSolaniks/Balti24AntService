from flask import Flask, request, render_template, render_template_string, send_file, redirect, url_for
import pandas as pd
import io
import os
from docx import Document
import mysql.connector

app = Flask(__name__)

EXCEL_PATH = "aizpildits_akts.xlsx"
OUTPUT_PATH = "aizpildits_akts_jauns.xlsx"

FIELD_DESCRIPTIONS = {
    "akta_id": "Akta numurs",
    "signature_date": "Parakstīšanas datums",
    "client_name": "Klienta nosaukums",
    "client_reg": "Klienta reģistrācijas numurs",
    "site_address": "Objekta adrese",
    "call_reason": "Izsaukuma iemesls",
    "executor_name": "Izpildītāja nosaukums",
    "executor_reg": "Izpildītāja reģistrācijas numurs",
    "job_type": "Darba veids",
    "start_time": "Darba sākuma laiks",
    "end_time": "Darba beigu laiks",
    "work_description": "Veiktā darba apraksts",
    "materials": "Izmantotie materiāli",
    "equipment_status": "Iekārtu stāvoklis",
    "worker_count": "Strādnieku skaits",
    "direct_costs": "Tiešās izmaksas",
    "vat": "PVN",
    "total_with_vat": "Kopā ar PVN",
    "client_signature": "Klienta paraksts",
    "executor_signature": "Izpildītāja paraksts"
}

HTML_FORM = '''
<!doctype html>
<title>Akta Ģenerators</title>
<h2>Ievadiet akta numuru (akta_id)</h2>
<form method=post enctype=multipart/form-data>
  <input type=text name=akta_id>
  <input type=submit value=Meklēt>
</form>
{% if error %}
  <p style="color:red">{{ error }}</p>
{% endif %}
{% if data_row %}
  <h3>Visi dati:</h3>
  <ul>
  {% for key, val in data_row.items() %}
    {% set clean_val = val if val and val|lower != 'nan' else '' %}
    <li><b>{{ field_descriptions.get(key, key) }}</b>: {{ clean_val }}</li>
  {% endfor %}
  </ul>
  <a href="/download/{{ akta_id }}">Lejupielādēt Word</a>
{% elif missing_keys %}
  <h3>Aizpildiet trūkstošos laukus:</h3>
  <form method=post action="/complete">
    {% for key in missing_keys %}
      <label>{{ field_descriptions[key] }}:</label><br>
      <input type=text name="{{ key }}"><br><br>
    {% endfor %}
    <input type=hidden name=akta_id value="{{ akta_id }}">
    <input type=submit value=Apstiprināt>
  </form>
{% endif %}
'''

def load_excel():
    if not os.path.exists(EXCEL_PATH):
        raise FileNotFoundError(f"Excel fails to load: {EXCEL_PATH}")
    df = pd.read_excel(EXCEL_PATH)
    df['akta_id'] = df['akta_id'].astype(str)
    return df

def save_excel(df):
    df.to_excel(OUTPUT_PATH, index=False)

@app.route('/', methods=['GET', 'POST'])
def index():
    akta_id = request.args.get('akta_id', '').strip()
    if request.method == 'POST':
        akta_id = request.form.get('akta_id', '').strip()

    if akta_id:
        df = load_excel()
        row = df[df['akta_id'] == akta_id]
        if row.empty:
            return render_template_string(HTML_FORM, error="Akts nav atrasts", missing_keys=None, data_row=None, akta_id=akta_id, field_descriptions=FIELD_DESCRIPTIONS)

        row_dict = row.iloc[0].to_dict()
        missing_keys = [k for k, v in row_dict.items() if pd.isna(v) or v == ""]
        if missing_keys:
            return render_template_string(HTML_FORM, missing_keys=missing_keys, data_row=None, akta_id=akta_id, error=None, field_descriptions=FIELD_DESCRIPTIONS)
        else:
            return render_template_string(HTML_FORM, missing_keys=None, data_row=row_dict, akta_id=akta_id, error=None, field_descriptions=FIELD_DESCRIPTIONS)

    return render_template_string(HTML_FORM, missing_keys=None, data_row=None, error=None, field_descriptions=FIELD_DESCRIPTIONS)

@app.route('/complete', methods=['POST'])
def complete():
    akta_id = request.form['akta_id'].strip()
    df = load_excel()
    for key in request.form:
        if key != 'akta_id':
            df.loc[df['akta_id'] == akta_id, key] = request.form[key]

    save_excel(df)
    row = df[df['akta_id'] == akta_id]
    if row.empty:
        return render_template_string(HTML_FORM, error="Datu saglabāšana neizdevās", missing_keys=None, data_row=None, akta_id=akta_id, field_descriptions=FIELD_DESCRIPTIONS)

    row_dict = row.iloc[0].to_dict()
    return render_template_string(HTML_FORM, missing_keys=None, data_row=row_dict, akta_id=akta_id, error=None, field_descriptions=FIELD_DESCRIPTIONS)

@app.route('/form')
def form_view():
    task_id = request.args.get('task_id')
    return render_template("form.html", task_id=task_id)


@app.route('/download/<akta_id>')
def download(akta_id):
    df = pd.read_excel(OUTPUT_PATH)
    df['akta_id'] = df['akta_id'].astype(str)
    row = df[df['akta_id'] == str(akta_id)]
    if row.empty:
        return "Akts nav atrasts."

    # Сохраняем информацию об акте в базу данных
    conn = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="balti24db"
    )
    cursor = conn.cursor()
    file_data = output.read()
    insert_query = """
        INSERT INTO acts (order_id, akta_id, file_name, file_data)
    VALUES (%s, %s, %s, %s)
    """
    order_id = akta_id  # Если akta_id и order_id совпадают — иначе уточни!
    file_name = f"akts_{akta_id}.docx"

    cursor.execute(insert_query, (order_id, akta_id, file_name, file_data))
    conn.commit()
    cursor.close()
    conn.close()

    row_dict = row.iloc[0].to_dict()

    doc = Document()
    doc.add_heading(f"Akts Nr. {akta_id}", level=1)

    for key, val in row_dict.items():
        description = FIELD_DESCRIPTIONS.get(key, key)
        clean_val = "" if pd.isna(val) or str(val).lower() == "nan" else str(val)
        doc.add_paragraph(f"{description}: {clean_val}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f"akts_{akta_id}.docx")

if __name__ == '__main__':
    app.run(debug=True)
